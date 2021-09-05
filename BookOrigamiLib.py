#-------------------------------------------------------------------------------
# Name:        module1
# Purpose:
#
# Author:      Mark
#
# Created:     25/12/2014
# Copyright:   (c) Mark 2014
# Licence:     <your licence>
#-------------------------------------------------------------------------------
from docx import Document
from docx.shared import Mm
import PythonMagick

def main():
    pass

if __name__ == '__main__':
    main()

def OpenAndInitialiseDocX(filename):
    document = Document("Book Origami Template.docx")
    section = document.sections[0]
    #document.add_paragraph(filename)
    document.add_picture(filename, width=Mm(50))
    return document

def CalculateAndWriteDocX(document, myImage):
    px = myImage.load()
    document.add_paragraph(str(myImage.size[1]) + "mm * " + str(myImage.size[0]) + " sheets")


    #from tkMessageBox import showinfo
    #showinfo("Image Mode", str(myImage.mode))
    #for each column of image
    imgFoldDimensions = [] #blank list to insert dimensions for whole image

    for col in range(1,myImage.size[0]+1):
        #for each pixel in column
        #Assume that first pixel is white - if first pixel is black, then Length for first section will be '0'
        state = "White";
        #begin at repeat=0 (ie, first white-black section)
        repeats = 0;
        count = 0
        #print "Sheet {0}".format(col)
##        document.add_paragraph()
##        p = document.add_paragraph("Sheet ")
##        p.add_run(str(col)).bold = True

        #pageFoldDimensions = [col,(0,0,0)]
        DimensionsStart = 0
        DimensionsLength = 0
        DimensionsEnd = 0
        Dimensions = []
        for row in range(1,myImage.size[1]+1):
            #get each pixel colour 255=white, 0=black
            colour = px[col-1,row-1];
            #showinfo("Pixel", str(col-1) + "," + str(row-1) + " = " + str(colour))
            #print colour
            if colour == 255: #white
                if state == "White": #previous pixel was white, this pixel was white, thus increment counter for this section
                    count += 1;
                else: #state == "black": #previous pixel was black, this pixel is white, thus increment repeats print previous counter, reset and change state.
                    repeats += 1
                    #print"Len: {1}mm End: {2}mm".format(col, count, row-1)
                    #append the remainder of the dimensions
##                    p.add_run("Len: ")
##                    p.add_run(str(count))
##                    p.add_run("mm End: ")
##                    p.add_run(str(row-1))
##                    p.add_run("mm")
                    DimensionsLength=str(count)
                    DimensionsEnd=str(row-1)
                    Dimensions.append((DimensionsStart, DimensionsLength, DimensionsEnd))
                    #imgFoldDimensions.append([DimensionsCol, (DimensionsStart, DimensionsLength, DimensionsEnd)])
                    count = 1
                    state = "White";
            else: #ie, colour = 0 = black
                if state == "Black": #previous pixel was black, this pixel was black, thus increment counter for this section
                    count += 1;
                else: #state == "white": #previous pixel was white, this pixel was black, thus print previous counter, reset and change state.
                    #if repeats = 0, then this is the 'start' height
                    if repeats == 0:
                        #print "Start: {1}mm".format(col, count) ,
##                        p = document.add_paragraph("Start: ")
##                        p.add_run(str(count))
##                        p.add_run("mm ")
                        DimensionsStart=str(count)
                    else: #this is the second or subsequent start dimension
                        #print "Repeat Start, {2}mm".format(col, count, row-1) ,
##                        p = document.add_paragraph("Rpt:   ")
##                        p.add_run(str(row-1))
##                        p.add_run("mm ")
                        DimensionsStart=str(row-1)
                    count = 1
                    state = "Black";
            if (row==(myImage.size[1]) and colour==0): #col ended with black # this should be an edge case, but still a bug - image should have black centred vertically if possible.
                DimensionsLength=str(count)
                DimensionsEnd=str(row)
                Dimensions.append((DimensionsStart, DimensionsLength, DimensionsEnd))
                #imgFoldDimensions.append([DimensionsCol, (DimensionsStart, DimensionsLength, DimensionsEnd)])

        imgFoldDimensions.append([col,Dimensions])
        state = "White"
        count = 0
        repeats = 0

    from tkMessageBox import askyesno
    keepallfolds = askyesno("Keep all folds?", "Do you want to keep all folds?")

    if not keepallfolds:
        #at this point, each page with folds is listed in imgFoldDimensions list
        #print imgFoldDimensions
        currentfold=1
        for page in range(1,len(imgFoldDimensions)+1):
            #imgFoldDimensions is a list of lists
            #imgFoldDimensions[page-1] is a list, containing page number[0] and then a list of tuples[1]
            #print imgFoldDimensions[page-1] # this prints that list
            # page # = imgFoldDimensions[page-1][0]
            # list of tuples of folds is imgFoldDimensions[page-1][1]

            currentpagefolds = imgFoldDimensions[page-1][1]
            numberofFolds = len(currentpagefolds)

            print numberofFolds
            if numberofFolds == 1:
                currentfold = 1
                Dims = imgFoldDimensions[page-1][1][0]
                print Dims
                PrintLine(document, page, Dims[0], Dims[1], Dims[2])
                #print str(page) + " " + str(Dims[0]) + " " + str(Dims[1]) + " " + str(Dims[2])
            else: #there is more than one section for this sheet
                if currentfold > numberofFolds:
                    currentfold = 1
                Dims = imgFoldDimensions[page-1][1][currentfold-1]
                PrintLine(document, page, Dims[0], Dims[1], Dims[2])
                print str(page) + " " + str(Dims[0]) + " " + str(Dims[1]) + " " + str(Dims[2])
                currentfold += 1
                continue
    else:
        #at this point, each page with folds is listed in imgFoldDimensions list
        #print imgFoldDimensions
        currentfold=1
        for page in range(1,len(imgFoldDimensions)+1):
            #imgFoldDimensions is a list of lists
            #imgFoldDimensions[page-1] is a list, containing page number[0] and then a list of tuples[1]
            #print imgFoldDimensions[page-1] # this prints that list
            # page # = imgFoldDimensions[page-1][0]
            # list of tuples of folds is imgFoldDimensions[page-1][1]

            currentpagefolds = imgFoldDimensions[page-1][1]
            numberofFolds = len(currentpagefolds)

            print numberofFolds
            for fold in currentpagefolds:
                print fold
                Dims = fold
                print Dims
                PrintLine(document, page, Dims[0], Dims[1], Dims[2])
            #if numberofFolds == 1:
            #    currentfold = 1
            #    Dims = imgFoldDimensions[page-1][1][0]
            #    print Dims
            #    PrintLine(document, page, Dims[0], Dims[1], Dims[2])
                #print str(page) + " " + str(Dims[0]) + " " + str(Dims[1]) + " " + str(Dims[2])
            #else: #there is more than one section for this sheet
            #    if currentfold > numberofFolds:
            #        currentfold = 1
            #    Dims = imgFoldDimensions[page-1][1][currentfold-1]
            #    PrintLine(document, page, Dims[0], Dims[1], Dims[2])
            #    print str(page) + " " + str(Dims[0]) + " " + str(Dims[1]) + " " + str(Dims[2])
            #    currentfold += 1
            #    continue


def PrintLine(document, sheet, start, length, end):
    #print (start)
    #print length
    #print end
    assert int(end) == int(start) + int(length)
    document.add_paragraph()
    p = document.add_paragraph("Sheet ")
    p.add_run(str(sheet)).bold = True
    p = document.add_paragraph("Start: ")
    p.add_run(str(start))
    p.add_run("mm ")
    #p.add_run("Len: ")
    #p.add_run(str(length))
    #p.add_run("mm ")
    p.add_run("End: ")
    p.add_run(str(end))
    p.add_run("mm")



def RemoveWhiteColumns(myImage):
    from PIL import Image
##    from tkMessageBox import showinfo
##    showinfo("Input Image Mode (before remove cols", myImage.mode)
    out = Image.new(myImage.mode, myImage.size, None)
    px = myImage.load()
    opixel = out.load()
    currentoutputcol = 0
    for column in range(0, myImage.size[0]): #for each column in input image
        sum = 255*myImage.size[1]
        for pixel in range(0,myImage.size[1]): #for each pixel in column
            sum -= px[column, pixel]
            #print px[column, pixel]
        if sum > 0: #ie, column not blank
##            print "col ", column, " not blank"
            #copy column to output
            for pixel in range(0,myImage.size[1]):
                opixel[currentoutputcol,pixel] = px[column, pixel]
            currentoutputcol += 1
    out = out.crop((0,0,currentoutputcol, pixel+1))
    return out
