#-------------------------------------------------------------------------------
# Name:        module1
# Purpose:
#
# Author:      Mark
#
# Created:     29/11/2014
# Copyright:   (c) Mark 2014
# Licence:     <your licence>
#-------------------------------------------------------------------------------

def main():
    pass


if __name__ == '__main__':
    main()


from Tkinter import *
import sys
#sys.argv[0] argument is script name
#sys.argv[1] is bitmap name.
print 'Number of arguments:', len(sys.argv), 'arguments.'
print 'Argument List:', str(sys.argv)

if len(sys.argv) == 2: #convert bitmap from argument
    filename = sys.argv[1]
else: #open GUI



    from tkFileDialog import askopenfilename
    Tk().withdraw()
    filename = askopenfilename()

print "Opening Image"
from PIL import Image;
myImage=Image.open(filename);
px = myImage.load();
print "Image Opened"

#what I'd like to be able to do is:
    #provide a standard bitmap
    #automatically convert to monochrome bitmap (thought bubble - choose monochrome from red, green or blue channels? provide preview?)
    #display monochrome bitmap for sanity check
    #get user input for height of book and number of sheets
    #automatically scale image appropriately for pixels per mm for height
    #automatically scale image appropriately for 1 pixel per sheet for width


from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
import time

print "Opening document"
time.sleep(2)
document = Document("Book Origami Template.docx")
section = document.sections[0]
document.add_paragraph(filename)
document.add_picture(filename, width=Mm(50))
print "Document opened"


#or each column of image
for col in range(1,myImage.size[0]+1):
    #for each pixel in colum
    #Assume that first pixel is white (else, doesn't make sense because means to leave top edge of page unfolded.
    state = "White";
    #begin at repeat=0 (ie, first white-black section)
    repeats = 0;
    count = 0
    #print "Sheet {0}".format(col)
    document.add_paragraph()
    p = document.add_paragraph("Sheet ")
    p.add_run(str(col)).bold = True

    for row in range(1,myImage.size[1]+1):
        #get first pixel colour 255=white, 0=black
        colour = px[col-1,row-1];
        if colour == 255: #white
            if state == "White": #previous pixel was white, this pixel was white, thus increment counter for this section
                count += 1;

            else: #state == "black": #previous pixel was black, this pixel was white, thus print previous counter, reset and change state.
                repeats += 1
                #print "Length: {1}mm End: {2}mm".format(col, count, row-1)
                p.add_run("Len: ")
                p.add_run(str(count))
                p.add_run("mm End: ")
                p.add_run(str(row-1))
                p.add_run("mm")


                count = 1
                state = "White";
            #print("Pixel", col, " ", row, " = white")
        else: #ie, colour = 0 = black
            if state == "Black": #previous pixel was black, this pixel was black, thus increment counter for this section
                count += 1;

            else: #state == "white": #previous pixel was white, this pixel was black, thus print previous counter, reset and change state.

                #if repeats = 0, then this is the 'start' height
                if repeats == 0:
                    #print "Start: {1}mm".format(col, count) ,
                    p = document.add_paragraph("Start: ")
                    p.add_run(str(count))
                    p.add_run("mm ")
                else:
                    #print "Repeat Start, {2}mm".format(col, count, row-1) ,
                    p = document.add_paragraph("Rpt:   ")
                    p.add_run(str(row-1))
                    p.add_run("mm ")
                #print "{0} count= {1}".format(state, count)
                count = 1
                state = "Black";
            #print("Pixel", col, " ", row, " = black")
   # print(state, ": count= ", count);
    #print "End of Sheet {0}".format(col)

    state = "White"
    count = 0
    repeats = 0

document.save(filename + ".docx")
print "Complete"